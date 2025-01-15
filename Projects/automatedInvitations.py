import os
from docx import Document

loadPath = os.path.join('Files', 'Invitations', 'InvitationTemplate.docx')
savePath = os.path.join('Files', 'Invitations\\')

names = ['Prof. Plum', 'Miss Scarlet', 'Col. Mustard', 'Al Sweigert', 'RoboCop']

for name in names:
    doc = Document(loadPath)
    doc.add_heading('It would be a pleasure to have the company of', 1)
    doc.add_heading(name, 2)
    doc.add_heading('at 11010 Memory Lane on the Evening of', 1)
    doc.add_heading('April 1st', 3)
    doc.add_heading('at 7 o\'clock', 1)
    doc.save(savePath + 'Invitation_' + name + '.docx')